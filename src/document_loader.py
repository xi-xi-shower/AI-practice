from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from PIL import Image
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = frozenset({".pdf", ".md", ".txt", ".docx", ".png"})


class DocumentLoadError(Exception):
    """Raised when a supported document cannot be parsed."""


@dataclass(frozen=True, slots=True)
class DocumentChunk:
    text: str
    source: str
    file_type: str
    metadata: dict[str, object]


# 将文档路径转换为绝对来源路径。
def _source(path: Path) -> str:
    return str(path.resolve())


# 按页提取 PDF 中的非空文本。
def _load_pdf(path: Path) -> list[DocumentChunk]:
    try:
        reader = PdfReader(path)
        chunks: list[DocumentChunk] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if not text or not text.strip():
                continue
            chunks.append(
                DocumentChunk(
                    text=text,
                    source=_source(path),
                    file_type="pdf",
                    metadata={"page_number": page_number},
                )
            )
        return chunks
    except Exception as exc:
        raise DocumentLoadError(f"Failed to load PDF '{path}': {exc}") from exc


# 使用 UTF-8 或 GB18030 编码读取文本文件。
def _read_markdown(path: Path) -> str:
    decode_error: UnicodeDecodeError | None = None
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            decode_error = exc
        except OSError as exc:
            raise DocumentLoadError(
                f"Failed to read Markdown file '{path}': {exc}"
            ) from exc

    raise DocumentLoadError(
        f"Failed to decode Markdown file '{path}' as UTF-8 or GB18030"
    ) from decode_error


# 将 Markdown 原文加载为单个文档块。
def _load_markdown(path: Path) -> list[DocumentChunk]:
    return [
        DocumentChunk(
            text=_read_markdown(path),
            source=_source(path),
            file_type="markdown",
            metadata={},
        )
    ]


# 将 TXT 原文加载为单个文档块。
def _load_text(path: Path) -> list[DocumentChunk]:
    return [
        DocumentChunk(
            text=_read_markdown(path),
            source=_source(path),
            file_type="text",
            metadata={},
        )
    ]


_ocr_engine = None


# 延迟创建并复用 RapidOCR 引擎。
def _get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is None:
        try:
            from rapidocr_onnxruntime import RapidOCR

            _ocr_engine = RapidOCR()
        except Exception as exc:
            raise DocumentLoadError(f"Failed to initialize OCR engine: {exc}") from exc
    return _ocr_engine


# 从 PNG 图片中提取 OCR 文本和图片元数据。
def _load_png(path: Path) -> list[DocumentChunk]:
    try:
        with Image.open(path) as image:
            width, height = image.size
            mode = image.mode
        result, _ = _get_ocr_engine()(str(path))
        if not result:
            return []
        text_lines = [item[1].strip() for item in result if len(item) > 1 and item[1].strip()]
        if not text_lines:
            return []
        return [
            DocumentChunk(
                text="\n".join(text_lines),
                source=_source(path),
                file_type="png",
                metadata={
                    "width": width,
                    "height": height,
                    "mode": mode,
                    "ocr_engine": "rapidocr-onnxruntime",
                },
            )
        ]
    except DocumentLoadError as exc:
        raise DocumentLoadError(f"Failed to load PNG '{path}': {exc}") from exc
    except Exception as exc:
        raise DocumentLoadError(f"Failed to load PNG '{path}': {exc}") from exc


# 提取 DOCX 中的非空段落和表格行。
def _load_docx(path: Path) -> list[DocumentChunk]:
    try:
        document = Document(path)
        chunks: list[DocumentChunk] = []

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            if not paragraph.text.strip():
                continue
            chunks.append(
                DocumentChunk(
                    text=paragraph.text,
                    source=_source(path),
                    file_type="docx",
                    metadata={
                        "content_type": "paragraph",
                        "paragraph_index": paragraph_index,
                    },
                )
            )

        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                cell_texts = [cell.text for cell in row.cells]
                if not any(text.strip() for text in cell_texts):
                    continue
                chunks.append(
                    DocumentChunk(
                        text="\t".join(cell_texts),
                        source=_source(path),
                        file_type="docx",
                        metadata={
                            "content_type": "table_row",
                            "table_index": table_index,
                            "row_index": row_index,
                        },
                    )
                )

        return chunks
    except Exception as exc:
        raise DocumentLoadError(f"Failed to load DOCX '{path}': {exc}") from exc


_LOADERS: dict[str, Callable[[Path], list[DocumentChunk]]] = {
    ".pdf": _load_pdf,
    ".md": _load_markdown,
    ".txt": _load_text,
    ".docx": _load_docx,
    ".png": _load_png,
}


# 根据文件扩展名加载单个受支持文档。
def load_document(
    path: str | Path,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> list[DocumentChunk]:
    """Load one supported document and return its non-empty content chunks."""
    document_path = Path(path)
    if not document_path.exists():
        raise FileNotFoundError(f"Document does not exist: '{document_path}'")
    if not document_path.is_file():
        raise IsADirectoryError(f"Document path is not a file: '{document_path}'")

    extension = document_path.suffix.lower()
    loader = _LOADERS.get(extension)
    if loader is None:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(
            f"Unsupported document type '{extension or '<none>'}'. "
            f"Supported types: {supported}"
        )
    return split_chunks(loader(document_path), chunk_size, chunk_overlap)


# 按稳定顺序查找目录中的受支持文档。
def _document_paths(directory: Path) -> Iterable[Path]:
    paths = (
        path
        for path in directory.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    return sorted(paths, key=lambda path: path.as_posix().casefold())


# 递归加载目录中的全部受支持文档。
def load_directory(
    path: str | Path,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> list[DocumentChunk]:
    """Recursively load all supported documents in a directory."""
    directory = Path(path)
    if not directory.exists():
        raise FileNotFoundError(f"Directory does not exist: '{directory}'")
    if not directory.is_dir():
        raise NotADirectoryError(f"Path is not a directory: '{directory}'")

    chunks: list[DocumentChunk] = []
    for document_path in _document_paths(directory):
        chunks.extend(load_document(document_path, chunk_size, chunk_overlap))
    return chunks


def _validate_split_parameters(chunk_size: int, chunk_overlap: int) -> None:
    """校验文本块长度和重叠长度是否合法。"""
    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than 0")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap must be greater than or equal to 0")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size")


def _find_split_end(text: str, start: int, limit: int, chunk_overlap: int) -> int:
    """在字符上限内查找最后一个优先切分边界。"""
    minimum_end = start + chunk_overlap + 1
    paragraph_ends = [
        index + 2
        for index in range(start, limit - 1)
        if text.startswith("\n\n", index) and index + 2 >= minimum_end
    ]
    if paragraph_ends:
        return paragraph_ends[-1]
    newline_ends = [
        index + 1
        for index in range(start, limit)
        if text[index] == "\n" and index + 1 >= minimum_end
    ]
    if newline_ends:
        return newline_ends[-1]
    sentence_ends = [
        index + 1
        for index in range(start, limit)
        if text[index] in "。！？.!?" and index + 1 >= minimum_end
    ]
    return sentence_ends[-1] if sentence_ends else limit


def split_chunks(
    chunks: Iterable[DocumentChunk],
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> list[DocumentChunk]:
    """按自然边界切分长文本块并保留来源信息。"""
    _validate_split_parameters(chunk_size, chunk_overlap)
    result: list[DocumentChunk] = []
    for chunk in chunks:
        if len(chunk.text) <= chunk_size:
            result.append(chunk)
            continue

        start = 0
        chunk_index = 0
        while start < len(chunk.text):
            end = _find_split_end(
                chunk.text,
                start,
                min(start + chunk_size, len(chunk.text)),
                chunk_overlap,
            )
            metadata = dict(chunk.metadata)
            metadata.update(
                {"chunk_index": chunk_index, "char_start": start, "char_end": end}
            )
            result.append(
                DocumentChunk(chunk.text[start:end], chunk.source, chunk.file_type, metadata)
            )
            if end == len(chunk.text):
                break
            start = end - chunk_overlap
            chunk_index += 1
    return result
