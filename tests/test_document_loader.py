from pathlib import Path

import pytest
from docx import Document

import src.document_loader as document_loader
from src.document_loader import (
    DocumentChunk,
    DocumentLoadError,
    load_directory,
    load_document,
    split_chunks,
)


# 验证 Markdown 原文和元数据能够完整保留。
def test_load_markdown_preserves_content_and_metadata(tmp_path: Path):
    markdown = tmp_path / "知识.md"
    content = "# 标题\n\n**保留** Markdown 标记。\n"
    markdown.write_text(content, encoding="utf-8")

    chunks = load_document(markdown)

    assert len(chunks) == 1
    assert chunks[0].text == content
    assert chunks[0].source == str(markdown.resolve())
    assert chunks[0].file_type == "markdown"
    assert chunks[0].metadata == {}


# 验证 Markdown 能回退使用 GB18030 编码。
def test_load_markdown_falls_back_to_gb18030(tmp_path: Path):
    markdown = tmp_path / "legacy.md"
    markdown.write_bytes("中文内容".encode("gb18030"))

    chunks = load_document(markdown)

    assert chunks[0].text == "中文内容"


# 验证 TXT 原文和元数据能够完整保留。
def test_load_text_preserves_content_and_metadata(tmp_path: Path):
    path = tmp_path / "sample.txt"
    content = "中文文本\n第二行\n"
    path.write_text(content, encoding="utf-8")

    chunks = load_document(path)

    assert len(chunks) == 1
    assert chunks[0].text == content
    assert chunks[0].file_type == "text"
    assert chunks[0].source == str(path.resolve())
    assert chunks[0].metadata == {}


# 验证 TXT 能回退使用 GB18030 编码。
def test_load_text_falls_back_to_gb18030(tmp_path: Path):
    path = tmp_path / "legacy.txt"
    path.write_bytes("旧版中文".encode("gb18030"))

    chunks = load_document(path)

    assert chunks[0].text == "旧版中文"


# 验证 PNG OCR 文本和图片元数据的加载结果。
def test_load_png_uses_ocr_and_preserves_image_metadata(monkeypatch, tmp_path: Path):
    from PIL import Image

    path = tmp_path / "sample.png"
    Image.new("RGB", (120, 80), "white").save(path)

    class FakeOCR:
        # 返回用于测试的固定 OCR 识别结果。
        def __call__(self, image_path):
            assert image_path == str(path)
            return [([[0, 0], [1, 1]], "第一行", 0.99), ([[0, 2], [1, 3]], "第二行", 0.98)], None

    monkeypatch.setattr(document_loader, "_ocr_engine", FakeOCR())

    chunks = load_document(path)

    assert len(chunks) == 1
    assert chunks[0].text == "第一行\n第二行"
    assert chunks[0].file_type == "png"
    assert chunks[0].metadata == {
        "width": 120,
        "height": 80,
        "mode": "RGB",
        "ocr_engine": "rapidocr-onnxruntime",
    }


# 验证没有识别文字的 PNG 会被跳过。
def test_load_png_skips_when_ocr_finds_no_text(monkeypatch, tmp_path: Path):
    from PIL import Image

    path = tmp_path / "empty.png"
    Image.new("RGB", (10, 10), "white").save(path)

    class FakeOCR:
        # 返回空 OCR 识别结果。
        def __call__(self, image_path):
            return [], None

    monkeypatch.setattr(document_loader, "_ocr_engine", FakeOCR())

    assert load_document(path) == []


# 验证 PNG OCR 异常会包含源文件路径。
def test_load_png_wraps_ocr_errors_with_file_path(monkeypatch, tmp_path: Path):
    from PIL import Image

    path = tmp_path / "failure.png"
    Image.new("RGB", (10, 10), "white").save(path)

    class FailingOCR:
        # 模拟 OCR 引擎调用失败。
        def __call__(self, image_path):
            raise RuntimeError("OCR failed")

    monkeypatch.setattr(document_loader, "_ocr_engine", FailingOCR())

    with pytest.raises(DocumentLoadError, match=r"Failed to load PNG .*failure\.png"):
        load_document(path)


# 验证损坏的 PNG 会抛出清晰异常。
def test_load_corrupt_png_raises_clear_error(tmp_path: Path):
    path = tmp_path / "corrupt.png"
    path.write_bytes(b"not a png")

    with pytest.raises(DocumentLoadError, match="Failed to load PNG"):
        load_document(path)


# 验证 DOCX 的非空段落和表格行能够被提取。
def test_load_docx_extracts_non_empty_paragraphs_and_table_rows(tmp_path: Path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("   ")
    document.add_paragraph("第三段")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "值"
    document.save(path)

    chunks = load_document(path)

    assert [chunk.text for chunk in chunks] == ["第一段", "第三段", "名称\t值"]
    assert chunks[0].metadata == {
        "content_type": "paragraph",
        "paragraph_index": 0,
    }
    assert chunks[1].metadata == {
        "content_type": "paragraph",
        "paragraph_index": 2,
    }
    assert chunks[2].metadata == {
        "content_type": "table_row",
        "table_index": 0,
        "row_index": 0,
    }
    assert all(chunk.file_type == "docx" for chunk in chunks)
    assert all(chunk.source == str(path.resolve()) for chunk in chunks)


# 验证 PDF 仅提取包含文本的页面。
def test_load_pdf_extracts_non_empty_pages(monkeypatch, tmp_path: Path):
    path = tmp_path / "sample.pdf"
    path.write_bytes(b"placeholder")

    class FakePage:
        # 保存测试页面的模拟文本。
        def __init__(self, text: str | None):
            self.text = text

        # 返回测试页面的模拟文本。
        def extract_text(self) -> str | None:
            return self.text

    class FakeReader:
        # 构造包含文本页和空白页的模拟 PDF。
        def __init__(self, file_path: Path):
            assert file_path == path
            self.pages = [FakePage("第一页"), FakePage("  "), FakePage(None), FakePage("第四页")]

    monkeypatch.setattr(document_loader, "PdfReader", FakeReader)

    chunks = load_document(path)

    assert [chunk.text for chunk in chunks] == ["第一页", "第四页"]
    assert [chunk.metadata for chunk in chunks] == [
        {"page_number": 1},
        {"page_number": 4},
    ]
    assert all(chunk.file_type == "pdf" for chunk in chunks)


# 验证目录加载支持递归过滤和稳定排序。
def test_load_directory_is_recursive_filtered_and_sorted(tmp_path: Path):
    nested = tmp_path / "nested"
    nested.mkdir()
    (tmp_path / "b.md").write_text("B", encoding="utf-8")
    (tmp_path / "a.md").write_text("A", encoding="utf-8")
    (nested / "c.md").write_text("C", encoding="utf-8")
    (tmp_path / "ignored.csv").write_text("ignored", encoding="utf-8")

    chunks = load_directory(tmp_path)

    assert [chunk.text for chunk in chunks] == ["A", "B", "C"]


# 验证单文件加载拒绝不支持的扩展名。
@pytest.mark.parametrize("filename", ["document.csv", "no_extension"])
def test_load_document_rejects_unsupported_types(tmp_path: Path, filename: str):
    path = tmp_path / filename
    path.write_text("content", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported document type"):
        load_document(path)


# 验证单文件加载拒绝不存在的路径。
def test_load_document_rejects_missing_file(tmp_path: Path):
    with pytest.raises(FileNotFoundError, match="Document does not exist"):
        load_document(tmp_path / "missing.pdf")


# 验证单文件加载拒绝目录路径。
def test_load_document_rejects_directory(tmp_path: Path):
    with pytest.raises(IsADirectoryError, match="not a file"):
        load_document(tmp_path)


# 验证目录加载能够区分缺失路径和文件路径。
def test_load_directory_validates_path(tmp_path: Path):
    file_path = tmp_path / "file.md"
    file_path.write_text("content", encoding="utf-8")

    with pytest.raises(FileNotFoundError, match="Directory does not exist"):
        load_directory(tmp_path / "missing")
    with pytest.raises(NotADirectoryError, match="not a directory"):
        load_directory(file_path)


# 验证损坏的 PDF 会抛出清晰异常。
def test_load_corrupt_pdf_raises_clear_error(tmp_path: Path):
    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"not a pdf")

    with pytest.raises(DocumentLoadError, match="Failed to load PDF"):
        load_document(path)


# 验证损坏的 DOCX 会抛出清晰异常。
def test_load_corrupt_docx_raises_clear_error(tmp_path: Path):
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a docx")

    with pytest.raises(DocumentLoadError, match="Failed to load DOCX"):
        load_document(path)


# 验证无法解码的 Markdown 会抛出清晰异常。
def test_load_undecodable_markdown_raises_clear_error(tmp_path: Path):
    path = tmp_path / "broken.md"
    path.write_bytes(b"\x81")

    with pytest.raises(DocumentLoadError, match="Failed to decode Markdown"):
        load_document(path)


def test_split_chunks_prefers_boundaries_and_preserves_overlap_metadata():
    chunk = DocumentChunk("alpha\n\nbeta\n\ngamma", "source", "text", {"page_number": 1})

    chunks = split_chunks([chunk], chunk_size=9, chunk_overlap=2)

    assert [item.text for item in chunks] == ["alpha\n\nb", "\nbeta\n\ng", "\ngamma"]
    assert [item.metadata for item in chunks] == [
        {"page_number": 1, "chunk_index": 0, "char_start": 0, "char_end": 8},
        {"page_number": 1, "chunk_index": 1, "char_start": 6, "char_end": 14},
        {"page_number": 1, "chunk_index": 2, "char_start": 12, "char_end": 18},
    ]


def test_split_chunks_prefers_newlines_then_sentence_endings_then_hard_cuts():
    newline_chunk = DocumentChunk("ab\ncd.efgh", "source", "text", {})
    hard_cut_chunk = DocumentChunk("abcdefgh", "source", "text", {})

    assert [item.text for item in split_chunks([newline_chunk], 5, 0)] == ["ab\n", "cd.", "efgh"]
    assert [item.text for item in split_chunks([hard_cut_chunk], 3, 0)] == ["abc", "def", "gh"]


def test_split_chunks_keeps_short_chunks_and_validates_parameters():
    chunk = DocumentChunk("short", "source", "text", {"existing": True})

    assert split_chunks([chunk], 5, 0) == [chunk]
    for chunk_size, chunk_overlap in ((0, 0), (1, -1), (5, 5)):
        with pytest.raises(ValueError):
            split_chunks([chunk], chunk_size, chunk_overlap)


def test_load_entrypoints_accept_custom_split_parameters(tmp_path: Path):
    (tmp_path / "a.md").write_text("abcdefgh", encoding="utf-8")
    nested = tmp_path / "nested"
    nested.mkdir()
    (nested / "b.txt").write_text("ijklmnop", encoding="utf-8")

    assert [item.text for item in load_document(tmp_path / "a.md", 3, 1)] == [
        "abc", "cde", "efg", "gh"
    ]
    assert [item.text for item in load_directory(tmp_path, 4, 0)] == [
        "abcd", "efgh", "ijkl", "mnop"
    ]
